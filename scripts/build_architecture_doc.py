from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = ROOT / "assets"
OUT = DOCS / "5月赛事_授权式协作记忆Skill_架构与Eval方案.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_font(run, size=None, bold=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    set_font(run, 18 if level == 1 else 14, True)
    return p


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, 10.5, True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2, 10.5)
    else:
        run = p.add_run(text)
        set_font(run, 10.5)
    p.paragraph_format.space_after = Pt(5)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run, 10.5)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_font(run, 10.5)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "D9EAF7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_font(run, 9.5, True)
    for row in rows:
        cells = table.add_row().cells
        for idx, val in enumerate(row):
            cells[idx].text = str(val)
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    set_font(run, 9)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_image(doc, filename, caption):
    path = ASSETS / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(8.7))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_font(r, 9, False)
    doc.add_paragraph()


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(8.5)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)


def build_doc():
    DOCS.mkdir(exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("5 月赛事：授权式协作记忆 Skill 架构与自动 Eval 工作台方案")
    set_font(run, 22, True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("面向用户授权记忆、反馈学习、工作流适应与持续协作质量提升")
    set_font(r, 12)
    doc.add_paragraph()

    add_heading(doc, "1. 方案目标", 1)
    add_para(
        doc,
        "本方案面向 5 月 AI 技能锦标赛主题：在用户授权下记住偏好、学习反馈、适应工作流，并在持续使用中提升协作质量。"
        "系统不是一个简单的长期记忆仓库，而是一个可复测、可审计、可对比、可自动评分的协作记忆 Skill。"
    )
    add_bullets(
        doc,
        [
            "从空白状态开始可复测，评委可以重放三轮对话并验证记忆变化。",
            "只提取长期偏好、场景规则、工作方法，过滤临时信息和不应保存信息。",
            "第二轮、第三轮任务能主动自然地应用记忆，减少用户重复说明。",
            "用户反馈变化时，旧记忆可以替换、降权、归档或删除。",
            "所有记忆的来源、使用原因、影响范围和版本变化都可查看。",
            "最终任务输出必须可直接交付，而不是只展示记忆机制。"
        ],
    )

    add_heading(doc, "2. 评分标准映射", 1)
    add_table(
        doc,
        ["评分维度", "权重", "核心考察点", "架构响应"],
        [
            ["可复测性", "10", "能否从空白状态开始，评委能否完全掌控并验证记忆。", "Reset、Replay、Trace、Snapshot、Eval Case 脚本。"],
            ["有效记忆提取", "20", "区分长期偏好、场景规则、工作方法，排除临时信息。", "候选提取、价值判断、分类器、证据与 scope 记录。"],
            ["记忆应用效果", "25", "第二次、第三次任务能主动自然应用记忆。", "记忆检索、自适应计划器、应用理由、前后输出对比。"],
            ["记忆更新与淘汰", "20", "偏好变化时旧记忆被替换、降权或归档。", "冲突检测、状态机、版本快照、回归检测。"],
            ["用户控制与透明度", "10", "记忆可查看、可解释、可干预。", "授权门禁、透明度面板、编辑/删除/拒绝保存。"],
            ["结果质量与可用性", "15", "最终任务输出可直接交付使用。", "任务执行器、质量评分、可交付结果模板。"],
        ],
        widths=[1.2, 0.5, 3.0, 4.0],
    )

    add_heading(doc, "3. 总体架构", 1)
    add_para(
        doc,
        "系统分为三条路径：主执行路径负责完成用户任务；记忆治理路径负责提取、分类、授权、更新和淘汰记忆；"
        "评估与看板路径负责自动评分、版本对比和趋势分析。"
    )
    add_image(doc, "01_architecture_overview.png", "图 1：授权式协作记忆与工作流适应 Skill 总体架构图（GPT image2.0 生成）")

    add_table(
        doc,
        ["模块", "职责", "关键产物"],
        [
            ["Consent Gate", "控制是否允许长期保存、更新或删除记忆。", "授权记录、拒绝保存记录、敏感信息拦截结果。"],
            ["Task Profiler", "识别任务类型、场景、目标、约束和是否需要调用记忆。", "task_profile、memory_query。"],
            ["Memory Extractor", "从用户输入、反馈、修改中提取候选记忆。", "memory_candidates。"],
            ["Memory Classifier", "区分长期偏好、场景规则、工作方法、项目上下文、临时信息。", "typed_memory_items。"],
            ["Memory Governance", "处理冲突、置信度、替换、降权、归档、删除。", "memory_patch、state_transition。"],
            ["Adaptive Planner", "把相关记忆转成当前任务的执行策略。", "execution_plan、applied_memory_reasons。"],
            ["Auto Eval Engine", "执行结构校验、规则评分、质量评分和加权汇总。", "score_breakdown、regression_alert。"],
            ["Eval Workbench", "图形化呈现三轮详情、版本对比、趋势图和记忆 diff。", "dashboard、report、export。"],
        ],
    )

    add_heading(doc, "4. 三轮持续协作 Eval 流程", 1)
    add_para(
        doc,
        "每个 eval case 采用三轮设计，确保评测覆盖“记忆形成、记忆应用、记忆更新与淘汰”的完整生命周期。"
    )
    add_image(doc, "02_three_round_eval_flow.png", "图 2：三轮持续协作 Eval 流程图（GPT image2.0 生成）")
    add_table(
        doc,
        ["轮次", "目标", "检查点"],
        [
            ["Round 1", "从空白状态提取偏好/工作方法，并请求授权保存。", "是否正确分类、是否过滤临时信息、是否生成 V1 快照。"],
            ["Round 2", "在同类任务中主动应用记忆，减少用户重复说明。", "是否检索相关记忆、是否自然应用、输出是否优于 V0/V1。"],
            ["Round 3", "用户偏好变化或提出反馈时，更新或淘汰旧记忆。", "是否检测冲突、替换/降权/归档是否合理、是否写入 V2 快照。"],
        ],
    )

    add_heading(doc, "5. 记忆生命周期与治理", 1)
    add_para(
        doc,
        "记忆治理的核心是“少记、准记、可解释、可淘汰”。系统不把所有用户信息都保存为长期记忆，而是先做价值判断、风险判断和授权确认。"
    )
    add_image(doc, "03_memory_lifecycle.png", "图 3：记忆生命周期与治理流程图（GPT image2.0 生成）")
    add_code_block(
        doc,
        """memory:
  id: mem_001
  type: communication_preference | workflow_rule | scene_rule | project_context
  content: "用户做架构方案时希望先分析，不要直接写代码"
  scope: global | project | task_type
  source: explicit_feedback | repeated_behavior | correction
  confidence: 0.90
  status: active | superseded | archived | deleted
  evidence:
    - "用户说：先不要开始写代码，需要分析架构方案"
  applies_when:
    - architecture_design
    - competition_skill_planning
  user_approved: true""",
    )

    add_heading(doc, "6. 自动 Eval 引擎设计", 1)
    add_para(
        doc,
        "自动 Eval 引擎由三层评分组成：结构合法性检查、确定性规则评分、LLM Judge 质量评分。"
        "最终按官方六维权重汇总，并额外执行版本回归检测。"
    )
    add_image(doc, "05_auto_eval_engine.png", "图 4：自动 Eval 引擎设计图（GPT image2.0 生成）")
    add_table(
        doc,
        ["评分层", "适合判断", "示例指标"],
        [
            ["结构校验", "可复测性与基础可运行", "reset 是否生效、schema 是否合法、trace/snapshot 是否存在。"],
            ["规则评分", "确定性行为", "未授权不写入、删除后不再应用、临时信息不长期保存。"],
            ["LLM Judge", "语义质量", "记忆提取是否准确、应用是否自然、结果是否可交付。"],
            ["回归检测", "进化是否正收益", "当前版本是否低于历史最佳，成本是否过高。"],
        ],
    )
    add_code_block(
        doc,
        """total_score =
  reproducibility * 0.10
  + memory_extraction * 0.20
  + memory_application * 0.25
  + update_and_decay * 0.20
  + transparency * 0.10
  + result_quality * 0.15""",
    )

    add_heading(doc, "7. 版本快照、对比与趋势分析", 1)
    add_para(
        doc,
        "每次记忆状态变化、每轮任务输出和每次评分结果都写入快照。快照用于 V0/V1/V2/V3 横向对比，也用于趋势图和回归检测。"
    )
    add_image(doc, "06_version_snapshot_trends.png", "图 5：版本快照、对比与趋势分析图（GPT image2.0 生成）")
    add_bullets(
        doc,
        [
            "V0 Baseline：无记忆状态下的任务输出和评分。",
            "V1：完成第一轮记忆提取与授权保存后的状态。",
            "V2：第二轮主动应用记忆后的状态。",
            "V3：第三轮根据用户反馈更新、降权或归档后的状态。",
            "Vn：持续使用中的当前版本，与历史最佳版本做回归对比。"
        ],
    )
    add_code_block(
        doc,
        """net_evolution_gain =
  current_total_score
  - previous_total_score
  - regression_penalty
  - cost_penalty""",
    )

    add_heading(doc, "8. 图形界面 Eval 工作台", 1)
    add_para(
        doc,
        "Eval 工作台用于让评委和开发者直观看到每个三轮测试的六维评分、记忆变化、版本差异和趋势。"
        "它也是透明度和可复测性的产品化展示层。"
    )
    add_image(doc, "04_eval_workbench_demo.png", "图 6：Eval 工作台 Demo 图（GPT image2.0 生成）")
    add_table(
        doc,
        ["页面/区域", "功能"],
        [
            ["Case Runner", "选择 eval case，一键 reset，运行 Round 1/2/3，导出报告。"],
            ["三轮详情", "展示每轮用户输入、使用记忆、记忆变更、输出质量和得分。"],
            ["Score Matrix", "按 case 展示六维评分和总分，支持筛选低分维度。"],
            ["Version Comparator", "左右对比 Vn 与 Vn-1、Vn 与 Best，展示输出 diff 和记忆 diff。"],
            ["Memory Timeline", "查看每条记忆从创建、应用、更新、归档到删除的生命周期。"],
            ["Trend Dashboard", "总分趋势、六维雷达图、token 成本、响应时间、记忆状态数量趋势。"],
        ],
    )

    add_heading(doc, "9. Eval Case 设计", 1)
    add_para(doc, "建议至少准备 10 个可复测 case，每个 case 都包含三轮输入、期望记忆、期望行为和评分规则。")
    add_table(
        doc,
        ["Case", "测试重点", "三轮设计"],
        [
            ["01 空白状态启动", "可复测性", "Reset -> 确认无隐藏记忆 -> 导出空快照。"],
            ["02 显式偏好提取", "有效记忆提取", "用户表达长期偏好 -> 授权保存 -> 检查结构化字段。"],
            ["03 临时信息过滤", "有效记忆提取", "用户说“这次用表格” -> 不写长期记忆 -> 会话内使用。"],
            ["04 工作方法提取", "有效记忆提取", "用户要求先读评分、再架构、再 eval -> 保存 workflow_rule。"],
            ["05 第二轮自然应用", "记忆应用效果", "同类任务中主动按 workflow_rule 输出，不要求用户重复说明。"],
            ["06 第三轮增强应用", "记忆应用效果", "进一步减少重复澄清，输出结构更贴合偏好。"],
            ["07 偏好冲突更新", "更新与淘汰", "用户从详细改成简版 -> 旧记忆降权或归档。"],
            ["08 删除控制", "用户控制与透明度", "用户要求忘记某条记忆 -> 后续不得再使用。"],
            ["09 透明度检查", "用户控制与透明度", "展示本轮使用了哪些记忆、为什么使用、来源是什么。"],
            ["10 最终可交付质量", "结果质量与可用性", "完成真实任务输出，检查能否直接用于参赛方案。"],
        ],
    )

    add_heading(doc, "10. 推荐落地顺序", 1)
    add_numbered(
        doc,
        [
            "先实现记忆 schema、reset/replay、snapshot，保证可复测。",
            "实现候选记忆提取、分类和授权门禁，优先解决有效记忆提取。",
            "实现检索与自适应计划器，让第二轮、第三轮能自然应用记忆。",
            "实现冲突检测、替换、降权、归档和删除，覆盖更新淘汰。",
            "实现自动 eval 引擎，先规则评分，再接入 LLM Judge。",
            "实现 Eval 工作台，重点展示三轮评分、版本对比、记忆时间线和趋势图。",
        ],
    )

    add_heading(doc, "11. 结论", 1)
    add_para(
        doc,
        "这个方案的竞争力在于把“记忆能力”变成可证明的协作质量提升：评委可以从空白状态重放三轮测试，看到记忆如何被提取、应用、更新和淘汰；"
        "也可以通过工作台比较 V0、V1、V2、V3 以及历史最佳版本，判断每次进化是否真的产生正收益。"
    )

    doc.add_page_break()
    add_heading(doc, "第二部分：项目计划", 1)
    add_para(
        doc,
        "最新协作结论：研发前期不建议按模块把两个人硬切开。这个项目的核心不确定性在于“哪条实现路线能在评分标准下拿到更高、更稳定的分数”，"
        "因此应采用 Eval-first、双路线端到端原型、单主线收敛的方式。这样耦合性最低，同时保留最大技术灵活性。"
    )

    add_heading(doc, "12. 协作策略", 1)
    add_table(
        doc,
        ["阶段", "协作方式", "原因", "产出"],
        [
            ["前期", "共同冻结 Eval 标准，不急于分模块开发。", "先确定裁判，再比较实现路线，避免主观争论。", "评分细则、10 个三轮 case、统一输出 schema。"],
            ["探索期", "两个人各自做端到端最小原型，而不是拆模块拼装。", "AI 辅助开发下，一个人打通完整闭环通常比多人拼模块更快。", "A/B 两条可运行路线。"],
            ["选择期", "用同一套 Eval 比较两条路线。", "按分数、稳定性、可优化性选主线。", "主线选择报告、失败 case 列表。"],
            ["收敛期", "一个人主攻主线，另一个人做评测、攻击、低分 case 分析和补强。", "减少代码冲突和接口等待，把第二个人变成质量杠杆。", "稳定版本、回归结果、工作台展示。"],
            ["封版期", "停止新增大功能，只修阻断问题和低分项。", "避免最后阶段发散，保证可演示、可复测。", "冻结版本、演示路径、导出报告。"],
        ],
        widths=[0.9, 2.2, 3.0, 2.8],
    )
    add_bullets(
        doc,
        [
            "原则一：前期并行探索，不并行拼装。",
            "原则二：中期 Eval 选主线，不靠讨论选方案。",
            "原则三：后期单线收敛，不继续发散。",
            "原则四：最终目标不是功能最多，而是在六维评分下稳定、可解释、可复测。"
        ],
    )

    add_heading(doc, "13. 时间计划", 1)
    add_para(
        doc,
        "交付日为 2026-06-16。研发应至少提前一周收口，最后一周留作视频、上报赛事、材料整理和突发问题 Buffer。"
    )
    add_table(
        doc,
        ["时间", "阶段", "核心目标", "验收标准"],
        [
            ["5/20-5/22", "Eval-first", "冻结六维评分、10 个三轮 case、snapshot schema、runner 输出协议。", "不用完整工作台，但必须有可执行的 Eval Contract。"],
            ["5/23-5/26", "双路线原型", "两个人各自打通一个最小闭环：提取、保存、应用、更新、快照。", "至少 1-2 个 case 能从空白状态跑通三轮。"],
            ["5/27", "统一 Eval 对比", "用同一套 eval 比较两条路线。", "输出总分、六维分、失败率、稳定性和可继续优化性。"],
            ["5/28", "选主线", "采用得分更高、风险更低、后续更容易封版的一条。", "确定主线代码和后续任务清单。"],
            ["5/29-6/03", "单主线强化", "补齐 10 个 case、自动评分、版本快照、回归检测、低分 case 修复。", "10 个 case 可自动跑并生成 JSON/Markdown 报告。"],
            ["6/04-6/06", "工作台与稳定性", "实现三轮详情、六维评分、版本对比、记忆时间线、趋势图。", "工作台能展示完整评测结果，主 demo 路径稳定。"],
            ["6/07", "研发初版冻结", "停止新增大功能，只接受阻断级 bug fix。", "可用、可演示、可复测的整体方案冻结。"],
            ["6/08", "内部验收", "完整回归，检查导出报告、截图和 demo 流程。", "进入材料 Buffer。"],
            ["6/09-6/15", "比赛材料 Buffer", "视频、讲解稿、截图、上报材料和最后修补。", "不再做架构级改动。"],
            ["6/16", "最终交付", "提交赛事材料。", "版本、报告、视频、说明材料齐备。"],
        ],
        widths=[1.15, 1.25, 4.0, 3.0],
    )

    add_heading(doc, "14. A/B 原型比较标准", 1)
    add_para(
        doc,
        "双路线开发不是为了长期维护两套代码，而是为了在早期用较低成本发现更高分、更稳定的实现路径。两条路线都必须遵守同一套 Eval Contract。"
    )
    add_table(
        doc,
        ["比较项", "判断方式", "采用倾向"],
        [
            ["总分", "10 个 case 跑完后的加权总分。", "总分更高者优先。"],
            ["六维短板", "观察是否存在明显低分维度。", "短板更少、修复路径更清晰者优先。"],
            ["稳定性", "同一 case 多次运行，比较结构一致性和失败率。", "输出更稳定者优先。"],
            ["记忆误存风险", "检查临时信息、敏感信息是否被错误长期保存。", "误存更少者优先。"],
            ["更新淘汰能力", "偏好冲突时是否替换、降权、归档旧记忆。", "状态转换更清楚者优先。"],
            ["实现复杂度", "后续是否容易补工作台、报告和回归检测。", "更容易封版者优先。"],
        ],
    )
    add_code_block(
        doc,
        """route_score =
  total_eval_score
  - failure_rate_penalty
  - instability_penalty
  - implementation_risk_penalty

adopt_route = route with higher route_score and clearer path to freeze""",
    )

    doc.add_page_break()
    add_heading(doc, "第三部分：第一步要做的 Eval 工作", 1)
    add_para(
        doc,
        "第一阶段不需要先搭完整图形化 Eval 工作台。应先搭 Eval 标准和最小可运行评测闭环；工作台只需要先冻结数据结构和页面草图。"
        "完整 dashboard 应在主线选定后再做，避免 UI 被早期实验实现绑死。"
    )

    add_heading(doc, "15. 第一阶段必须完成", 1)
    add_table(
        doc,
        ["事项", "内容", "完成标准"],
        [
            ["六维评分标准", "可复测性、有效记忆提取、记忆应用效果、记忆更新与淘汰、用户控制与透明度、结果质量与可用性。", "每个维度有权重、评分点、扣分点和 0-满分分档。"],
            ["10 个三轮 Eval Case", "每个 case 包含 Round 1 记忆形成、Round 2 记忆应用、Round 3 更新淘汰。", "case 可从空白状态自动回放。"],
            ["Eval Contract", "统一定义输入、输出、记忆前后状态、快照、评分结果。", "两条实现路线都能输出同一格式。"],
            ["Runner 协议", "定义如何 reset、如何运行 case、如何收集结果。", "可以用 CLI 跑出 eval_report.json。"],
            ["最小报告", "JSON 与 Markdown 报告。", "能看出谁得分高、低分项在哪里、是否有回归。"],
            ["工作台数据草图", "定义后续 dashboard 要读取哪些数据。", "暂不要求完整 UI，只保证数据产物可支撑工作台。"],
        ],
        widths=[1.5, 4.1, 3.4],
    )

    add_heading(doc, "16. 第一阶段暂不做", 1)
    add_bullets(
        doc,
        [
            "不先做完整图形化工作台。",
            "不先做复杂趋势图交互。",
            "不先做多用户、多项目管理。",
            "不先做精美 dashboard 和视频级展示页。",
            "不为了 UI 设计提前绑定某条实现路线。"
        ],
    )

    add_heading(doc, "17. Eval Contract 建议", 1)
    add_para(
        doc,
        "所有实现路线必须输出统一结构，自动评测和后续工作台只依赖这个结构，不依赖内部实现。"
    )
    add_code_block(
        doc,
        """eval_result:
  run_id: run_2026_05_22_001
  route_id: route_a | route_b | main
  case_id: workflow_adaptation_001
  round: 1 | 2 | 3
  task_input: ...
  task_output: ...
  memory_before:
    active: []
    superseded: []
    archived: []
  memory_after:
    active: []
    superseded: []
    archived: []
  applied_memories:
    - memory_id: mem_001
      reason: ...
  memory_changes:
    added: []
    updated: []
    archived: []
    deleted: []
  trace:
    decisions: []
    tool_calls: []
    errors: []
  scores:
    reproducibility: 0
    memory_extraction: 0
    memory_application: 0
    update_and_decay: 0
    transparency: 0
    result_quality: 0
    total: 0
  judge_reasons:
    - dimension: memory_application
      reason: ...""",
    )

    add_heading(doc, "18. 六维评分细则", 1)
    add_table(
        doc,
        ["维度", "满分", "高分标准", "主要扣分点"],
        [
            ["可复测性", "10", "能 reset 到空白状态；case 可自动回放；trace、snapshot、report 完整。", "隐藏状态、无法重放、快照缺失、输出格式不稳定。"],
            ["有效记忆提取", "20", "能区分长期偏好、场景规则、工作方法、项目上下文、临时信息；有证据和 scope。", "关键词乱存、临时信息长期化、缺少证据、scope 错误。"],
            ["记忆应用效果", "25", "Round 2/3 主动自然应用记忆，减少重复说明，输出更贴合用户。", "后续不用记忆、应用生硬、套错记忆、输出无提升。"],
            ["记忆更新与淘汰", "20", "能检测冲突并替换、降权、归档或删除旧记忆；状态变化可追溯。", "只追加不淘汰、冲突不处理、删除后仍应用、版本不可追溯。"],
            ["用户控制与透明度", "10", "保存前授权；可查看、解释、修改、删除；说明本轮为何使用某条记忆。", "黑箱保存、未授权写入、无法解释来源、不可干预。"],
            ["结果质量与可用性", "15", "最终任务结果完整、准确、可直接交付，并体现记忆带来的个性化提升。", "只有机制 demo，任务结果不可用；内容空泛；个性化不明显。"],
        ],
        widths=[1.25, 0.6, 4.0, 3.4],
    )

    add_heading(doc, "19. 评分分档建议", 1)
    add_table(
        doc,
        ["维度", "低分", "中分", "高分"],
        [
            ["可复测性", "无法从空白状态稳定复现。", "能 reset 和输出报告，但 trace/snapshot 不完整。", "评委可完整控制初始状态、输入脚本、记忆快照和报告。"],
            ["有效记忆提取", "只做关键词保存，临时信息也乱记。", "能提取偏好，但分类和证据粗糙。", "能结构化区分类型、scope、证据、置信度，并过滤不应保存内容。"],
            ["记忆应用效果", "后续任务基本不用记忆。", "偶尔引用记忆，但应用不稳定或不自然。", "第二、第三轮都能自然应用，输出明显更贴合用户。"],
            ["记忆更新与淘汰", "新旧偏好简单堆叠。", "能发现冲突，但处理不稳定。", "有替换、降权、归档、删除、版本和回滚/追溯机制。"],
            ["用户控制与透明度", "用户不知道记了什么。", "能查看记忆，但解释和干预弱。", "能查看来源、使用原因、影响范围，并可修改、删除、拒绝保存。"],
            ["结果质量与可用性", "只有记忆机制，结果不可交付。", "结果基本可用，但记忆提升不明显。", "结果可直接交付，且记忆显著提升协作效率。"],
        ],
    )

    add_heading(doc, "20. 最小可运行 Eval 目录", 1)
    add_code_block(
        doc,
        """eval/
  cases/
    case_01_blank_state.yaml
    case_02_explicit_preference.yaml
    case_03_temporary_info_filter.yaml
    ...
  rubrics/
    may_memory_score.yaml
  outputs/
    run_YYYYMMDD_HHMM/
      eval_report.json
      eval_report.md
      snapshots/
        case_01_round_1.json
        case_01_round_2.json
        case_01_round_3.json
      comparisons/
        route_a_vs_route_b.json
        version_trend.json""",
    )

    add_heading(doc, "21. 第一阶段的成功标准", 1)
    add_bullets(
        doc,
        [
            "两个人的不同实现可以用同一个 runner 跑。",
            "同一个 case 可以从空白状态重放三轮。",
            "每轮都有 memory_before、memory_after、applied_memories、memory_changes 和 scores。",
            "eval_report.json 能直接回答哪条路线更好。",
            "Markdown 报告能让人快速看到六维分、低分 case 和失败原因。",
            "工作台暂时不完整，但后续可以直接读取这些 eval 产物。"
        ],
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
