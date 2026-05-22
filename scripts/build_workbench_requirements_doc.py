from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = DOCS / "Eval工作台_需求与研发说明.docx"


def set_font(run, size=None, bold=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_font(r, 18 if level == 1 else 14, True)


def para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, 10.5)
    p.paragraph_format.space_after = Pt(5)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r, 10.5)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        shade(c, "D9EAF7")
        for p in c.paragraphs:
            for r in p.runs:
                set_font(r, 9.5, True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    set_font(r, 9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(8.5)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)


def build():
    DOCS.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Eval 工作台需求与研发说明")
    set_font(r, 22, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("面向授权式协作记忆 Skill 的 Case 评测、记忆状态跃迁与模块消融对比")
    set_font(r, 12)

    heading(doc, "1. 产品定位", 1)
    para(
        doc,
        "Eval 工作台用于验证 5 月赛事 Skill 是否能在用户授权下提取有效记忆、自然应用记忆、更新或淘汰旧记忆，并最终降低用户协作成本。"
        "工作台不以“Skill 版本 S1/S2/S3”为主叙事，而以历史执行 Case、三轮 Round、Memory 状态跃迁和模块 ON/OFF 消融对比为核心。"
    )
    bullets(
        doc,
        [
            "让评委一眼看到：跑过哪些 Case，每个 Case 三轮发生了什么。",
            "让研发知道：每轮 Skill 做了哪些动作，记忆是新增、应用、更新、删除还是淘汰。",
            "让评分可解释：Round 不显示综合总分，Case 三轮完成后才汇总六维总分。",
            "让模块价值可证明：通过 OFF vs ON 展示关键模块带来的差异。",
            "让持续协作质量可见：通过 M0/M1/M2 的用户操作费力度下降展示增益。"
        ],
    )

    heading(doc, "2. 核心概念", 1)
    table(
        doc,
        ["概念", "定义", "是否打总分", "展示方式"],
        [
            ["Eval Case", "一个完整三轮测试场景，覆盖记忆形成、应用、更新/淘汰。", "是", "Case 列表和 Case 详情页。"],
            ["Round", "Case 内的一轮对话执行步骤。", "否", "只展示动作、增益、贡献维度和依据。"],
            ["Memory State", "用户记忆状态快照，如 M0/M1/M2。", "否", "展示状态跃迁和用户费力度。"],
            ["Capability Module", "Skill 的可开关能力，如冲突处理、临时信息过滤、删除控制。", "否", "OFF vs ON 消融对比。"],
            ["Case Score", "三轮完成后的六维加权评分。", "是", "Case 底部总分和六维分。"],
            ["Config Average", "当前模块配置跑全部 Case 后的平均分。", "是", "顶部概览指标。"],
        ],
        widths=[1.4, 3.7, 1.1, 3.0],
    )

    heading(doc, "3. 页面信息架构", 1)
    para(doc, "工作台建议采用单页结构，降低认知负担。默认只展示核心亮点，细节通过展开项和气泡查看。")
    table(
        doc,
        ["区域", "内容", "交互", "研发说明"],
        [
            ["顶部概览", "评测对象、历史 Case 数、当前模块配置平均分、核心增益。", "无或筛选。", "不要显示 S1/S2/S3 版本叙事。"],
            ["历史 Case 列表", "Case 名称、目标能力、总分、核心记忆动作标签。", "点击切换 Case。", "列表只保留摘要，避免大段文本。"],
            ["Case 详情头部", "Case 标题、目标、核心亮点一句话。", "无。", "一句话说明该 Case 证明了什么。"],
            ["三轮 Round 卡片", "Round 1/2/3 顺序排开，展示用户输入、Skill 动作、记忆动作、状态变化、本轮增益。", "展开依据。", "Round 不显示综合分。"],
            ["Memory 状态跃迁", "M0/M1/M2、状态说明、用户操作费力度高/中/低。", "费力度气泡；评分规则展开。", "用于展示持续使用中协作成本下降。"],
            ["模块消融对比", "某个模块 OFF 会怎样，ON 后改善什么。", "随 Case 切换。", "用于证明模块价值。"],
            ["Case 六维评分", "六维分和 Case 总分。", "无或 hover 解释。", "只有 Case 级有完整总分。"],
        ],
    )

    heading(doc, "4. Round 展示规范", 1)
    para(doc, "Round 是执行节点，不是完整评测对象。每个 Round 只展示本轮验证了哪些能力和产生了什么增益。")
    table(
        doc,
        ["字段", "示例", "是否默认显示"],
        [
            ["Round 名称", "Round 1：记忆形成", "是"],
            ["用户输入", "先不要写代码，按照评分标准分析架构方案。", "是，单行截断"],
            ["Skill 动作", "识别长期工作流偏好。", "是"],
            ["记忆动作", "新增 workflow_rule。", "是"],
            ["状态变化", "M0 空白 → M1 偏好已保存", "是"],
            ["本轮增益", "后续无需重复说明“先分析再实现”。", "是"],
            ["贡献维度", "有效记忆提取、用户控制、可复测性", "否，展开后显示"],
            ["判断依据", "证据来自用户显式表达，scope 为 architecture_design。", "否，展开后显示"],
        ],
    )
    para(doc, "Round 卡片中的记忆动作需要用醒目标签展示，推荐动作包括：新增、应用、更新、降权、归档、删除、未保存、避免误用、不再使用。")

    heading(doc, "5. Memory 状态跃迁与用户费力度", 1)
    para(
        doc,
        "Memory State 不打官方总分，但需要展示一个辅助指标：用户操作费力度。它用于证明记忆状态向前跃迁后，用户达到同样结果所需的重复说明、澄清和修正成本下降。"
    )
    table(
        doc,
        ["等级", "判定标准", "展示方式"],
        [
            ["高", "用户必须完整说明偏好/流程，或多次修正才能达到目标。", "红色/高费力度气泡。"],
            ["中", "系统记住部分偏好，但用户仍需补充条件或纠正边界。", "橙色/中费力度气泡。"],
            ["低", "用户只需给任务或说明例外，系统能自然应用记忆并减少澄清。", "绿色/低费力度气泡。"],
        ],
    )
    para(doc, "费力度气泡点击或悬停后，应展示该等级的判定依据。下方提供“展开费力度评分规则”，说明计算参考。")
    bullets(
        doc,
        [
            "重复说明次数：用户是否还要重复同一偏好。",
            "澄清轮数：系统是否还要反复问用户。",
            "修正次数：用户是否还要指出“不是这样”。",
            "指令长度：用户需要多长的输入才能达到目标。",
            "误用次数：旧记忆或错误记忆是否造成后续偏差。"
        ],
    )
    para(doc, "用户费力度采用 0-100 分，分数越高表示用户越费力，分数越低表示协作越省力。它不是官方六维总分的一部分，而是用于解释 Memory State 从 M0 到 M1/M2 后是否真的降低用户协作成本。")
    table(
        doc,
        ["子指标", "权重", "含义", "0 分", "100 分"],
        [
            ["重复说明成本", "30%", "用户是否需要重复讲已经表达过的偏好、流程或约束。", "无需重复说明，系统自动应用。", "需要完整重新说明。"],
            ["修正成本", "25%", "用户是否需要纠正模型输出方向、格式或偏好误用。", "无需修正。", "输出不可用，需要重新开始。"],
            ["澄清成本", "15%", "模型是否因为缺少记忆或不确定边界而反复询问。", "无需澄清。", "无法判断偏好，卡住不执行。"],
            ["输入复杂度", "20%", "用户为了达到目标需要写多长、多复杂的指令。", "一句短指令即可触发正确行为。", "需要完整 prompt 模板。"],
            ["记忆误用成本", "10%", "错误、过期、冲突或临时记忆是否造成干扰。", "无误用。", "错误记忆导致任务失败。"],
        ],
        widths=[1.5, 0.7, 3.1, 2.0, 2.0],
    )
    code(
        doc,
        """user_effort_score =
  repeated_instruction_score * 0.30
  + correction_score * 0.25
  + clarification_score * 0.15
  + input_complexity_score * 0.20
  + memory_misuse_score * 0.10

level =
  low    if score <= 30
  medium if 31 <= score <= 65
  high   if score >= 66""",
    )
    table(
        doc,
        ["总费力度分数", "UI 等级", "解释"],
        [
            ["0-30", "低", "用户只需给任务或说明例外，系统能自然适配。"],
            ["31-65", "中", "系统记住部分信息，但用户仍需补充条件或修正边界。"],
            ["66-100", "高", "用户需要重复说明、反复纠正或完整重写要求。"],
        ],
    )
    table(
        doc,
        ["Memory 状态", "重复说明", "修正", "澄清", "输入复杂度", "误用", "总费力度", "含义"],
        [
            ["M0 空白", "90", "40", "30", "80", "0", "59", "用户必须详细说明流程和偏好。"],
            ["M1 偏好已保存", "10", "15", "10", "20", "0", "12", "用户只需说任务，系统自动应用偏好。"],
            ["M2 条件化更新", "5", "5", "10", "15", "0", "7", "用户只需说明例外，系统能处理冲突。"],
        ],
    )
    para(doc, "状态跃迁重点看趋势：M1 应低于 M0；M2 应低于或至少不高于 M1。如果 M2 高于 M1，说明更新后的记忆可能带来额外复杂度或误用，需要回归检查。")
    code(
        doc,
        """"effort_signals": {
  "repeated_instruction_count": 1,
  "correction_count": 0,
  "clarification_turns": 1,
  "user_input_tokens": 42,
  "memory_misuse_count": 0
},
"effort_score": {
  "repeated_instruction": 75,
  "correction": 0,
  "clarification": 25,
  "input_complexity": 60,
  "memory_misuse": 0,
  "total": 34.5,
  "level": "medium"
}""",
    )

    heading(doc, "6. 模块消融对比", 1)
    para(doc, "工作台需要避免用 Skill 版本号讲故事。模块价值通过同一个 Case 下的 OFF vs ON 对比体现。")
    table(
        doc,
        ["模块", "OFF 典型表现", "ON 后改善", "关联评分维度"],
        [
            ["冲突处理模块", "新旧偏好同时 active，后续输出摇摆。", "旧规则降权或归档，新规则条件化生效。", "记忆更新与淘汰、记忆应用效果。"],
            ["临时信息过滤模块", "把“这次用表格”误存为长期偏好。", "仅会话内使用，不污染长期画像。", "有效记忆提取、用户控制。"],
            ["删除控制模块", "UI 显示删除，但检索层仍命中旧记忆。", "deleted 记忆被检索层过滤，后续不再使用。", "用户控制与透明度、可复测性。"],
            ["授权门禁模块", "未授权保存长期记忆。", "保存前请求授权，并记录来源与范围。", "用户控制与透明度。"],
            ["记忆应用模块", "后续任务不主动使用已保存偏好。", "Round 2/3 自然应用相关记忆，减少重复说明。", "记忆应用效果、结果质量。"],
        ],
    )

    heading(doc, "7. Case 级评分逻辑", 1)
    para(doc, "只有 Case 三轮完成后才产生完整六维总分。Round 只贡献局部证据，不单独显示综合总分。")
    table(
        doc,
        ["评分维度", "权重", "主要证据来源"],
        [
            ["可复测性", "10", "Reset、三轮 replay、trace、snapshot、状态变化可验证。"],
            ["有效记忆提取", "20", "Round 1 的记忆提取、分类、scope、evidence 和临时信息过滤。"],
            ["记忆应用效果", "25", "Round 2/3 是否自然应用记忆，是否减少用户重复说明。"],
            ["记忆更新与淘汰", "20", "Round 3 是否处理冲突、替换、降权、归档、删除。"],
            ["用户控制与透明度", "10", "授权、可解释、可查看、可删除、使用原因说明。"],
            ["结果质量与可用性", "15", "最终输出是否可交付，是否体现记忆带来的质量提升。"],
        ],
    )
    code(
        doc,
        """case_score =
  reproducibility
  + memory_extraction
  + memory_application
  + update_and_decay
  + transparency
  + result_quality

config_average =
  average(case_score for all completed eval cases under current module configuration)""",
    )

    heading(doc, "8. Eval Case 剧本设计", 1)
    para(
        doc,
        "每个 Case 以“三轮主流程 + 删除后复测”组织。删除后复测是控制验证步骤，不算 Round 4；它用于证明用户删除记忆后，检索层和执行层都真正停止使用该记忆。"
    )
    table(
        doc,
        ["阶段", "是否计入三轮", "目标", "工作台展示重点"],
        [
            ["清空记忆：reset memory", "控制步骤", "确保从空白状态开始，保证可复测。", "显示 M0 空白状态。"],
            ["首次任务：无偏好的普通任务", "Round 1 前半", "观察无记忆 baseline。", "显示未应用记忆。"],
            ["用户反馈：明确偏好和工作方法", "Round 1 后半", "生成有效记忆 M1。", "显示新增记忆、证据、scope、授权。"],
            ["查看记忆", "控制步骤", "验证记忆可见、可解释。", "显示 Memory Store 当前内容。"],
            ["再次任务：相似但不同任务", "Round 2", "观察是否主动应用 M1。", "显示应用记忆和用户费力度下降。"],
            ["偏好变化：推翻或缩小旧偏好范围", "Round 3 前半", "生成 M2，处理冲突或缩小适用范围。", "显示更新、降权、归档或条件化。"],
            ["第三次任务", "Round 3 后半", "观察新规则是否生效，旧规则是否停止影响。", "显示 M2 生效、旧记忆不再主导。"],
            ["删除后复测", "控制步骤", "删除某条记忆，再执行任务验证不再使用。", "显示 deleted 状态和不再命中。"],
        ],
        widths=[2.0, 1.1, 3.0, 3.0],
    )
    table(
        doc,
        ["Case", "场景", "人物", "核心模块", "主要验证点"],
        [
            ["C01", "生活：家庭旅行规划", "用户是成年子女，要给父亲和孩子安排周末行程。", "记忆提取、应用、条件化更新、删除控制", "长期偏好与场景规则能否降低反复说明。"],
            ["C02", "工作：项目周报与决策材料", "用户是项目经理，需要给老板和跨部门团队同步进展。", "工作方法记忆、格式偏好、冲突处理", "同类但不同任务中是否自然应用工作流。"],
            ["C03", "学习：考试复习计划", "用户是学生，需要规划不同科目的复习。", "学习偏好、时间策略、偏好缩小适用范围", "偏好变化后是否不再机械套用旧计划。"],
            ["C04", "研究：文献综述与研究设计", "用户是研究者，需要做文献梳理和研究问题设计。", "研究方法偏好、引用约束、模式切换", "从综述模式切到假设生成模式时，旧规则是否降权。"],
        ],
    )

    heading(doc, "8.1 Case C01：生活场景 - 家庭旅行规划", 2)
    table(
        doc,
        ["阶段", "评委输入/操作", "期望 Skill 行为", "记忆动作/评估点"],
        [
            ["Reset", "请执行 reset memory。", "清空记忆，展示 M0。", "可复测性。"],
            ["首次任务", "帮我安排北京周末 2 天亲子旅行。", "给出普通行程，不假设偏好。", "无记忆应用。"],
            ["反馈", "以后家庭出行请记住：父亲膝盖不好，步行要少；孩子喜欢自然和动物；我不喜欢人挤人的网红点。", "提取 family_travel_rule，并请求/记录授权。", "新增 M1：慢节奏、少步行、自然动物、避开高拥挤。"],
            ["查看记忆", "展示当前记忆。", "列出记忆内容、来源、scope、状态。", "透明度。"],
            ["再次任务", "帮我安排杭州 3 天家庭行程。", "主动应用 M1，少步行、安排休息、自然/动物、避开拥挤。", "用户费力度下降，记忆应用效果。"],
            ["偏好变化", "这次父亲不去，只有我和孩子；避开网红点这条保留，但少步行不适用。", "将“少步行”缩小到父亲同行场景。", "M1 条件化更新为 M2，旧规则降权。"],
            ["第三次任务", "帮我安排上海 1 天亲子自然路线。", "不再强制少步行，但仍避开拥挤、偏自然。", "新规则生效，旧规则停止泛化影响。"],
            ["删除复测", "删除‘孩子喜欢动物’这条记忆，再安排南京半日游。", "删除后不再主动安排动物园/动物主题。", "deleted 记忆不再命中。"],
        ],
    )

    heading(doc, "8.2 Case C02：工作场景 - 项目周报与决策材料", 2)
    table(
        doc,
        ["阶段", "评委输入/操作", "期望 Skill 行为", "记忆动作/评估点"],
        [
            ["Reset", "请执行 reset memory。", "清空记忆，展示 M0。", "可复测性。"],
            ["首次任务", "根据这些进展写一份项目周报。", "输出普通周报。", "无偏好 baseline。"],
            ["反馈", "以后写给老板的项目材料，请先给 3 条结论，再用表格列风险、负责人和下一步。", "提取 work_report_method。", "新增 M1：老板材料先结论、风险表格化。"],
            ["查看记忆", "展示当前记忆。", "展示来源、适用对象为老板/管理层材料。", "scope 准确性。"],
            ["再次任务", "把这段研发进展整理成给老板看的同步材料。", "自动先给 3 条结论，再列风险/负责人/下一步表格。", "应用 M1，减少重复说明。"],
            ["偏好变化", "跨部门同步不要那么管理层风格，风险表只用于老板材料。", "缩小 M1 的适用范围。", "M2：老板材料保留，跨部门材料不强制风险表。"],
            ["第三次任务", "写一份给设计、研发、运营的跨部门同步。", "不再套老板材料模板，改为协作事项和依赖。", "旧规则停止影响不适用场景。"],
            ["删除复测", "删除‘先给 3 条结论’这条记忆，再写老板材料。", "不再强制 3 条结论结构。", "删除控制、透明度。"],
        ],
    )

    heading(doc, "8.3 Case C03：学习场景 - 考试复习计划", 2)
    table(
        doc,
        ["阶段", "评委输入/操作", "期望 Skill 行为", "记忆动作/评估点"],
        [
            ["Reset", "请执行 reset memory。", "清空记忆，展示 M0。", "可复测性。"],
            ["首次任务", "帮我做一个 7 天英语复习计划。", "给出普通复习计划。", "无偏好 baseline。"],
            ["反馈", "以后学习计划请按 25 分钟番茄钟安排；我喜欢先看例题再讲知识点；每天最后要有 5 道自测题。", "提取 learning_preference。", "新增 M1：番茄钟、例题先行、每日自测。"],
            ["查看记忆", "展示当前记忆。", "展示偏好类型、证据、适用范围。", "透明度。"],
            ["再次任务", "帮我做一个 5 天高数复习计划。", "主动应用 M1：番茄钟、例题先行、自测。", "用户输入更短，费力度下降。"],
            ["偏好变化", "考试只剩两天了，番茄钟不用了，改成按高频考点冲刺；但例题先行还保留。", "将番茄钟降权/归档，保留例题先行。", "M2：冲刺模式优先。"],
            ["第三次任务", "帮我做 2 天线性代数冲刺。", "不再机械套 25 分钟番茄钟，改为高频考点冲刺，并保留例题先行。", "更新淘汰生效。"],
            ["删除复测", "删除‘每天 5 道自测题’这条记忆，再做物理复习计划。", "不再强制每日 5 题。", "deleted 记忆不再使用。"],
        ],
    )

    heading(doc, "8.4 Case C04：研究场景 - 文献综述与研究设计", 2)
    table(
        doc,
        ["阶段", "评委输入/操作", "期望 Skill 行为", "记忆动作/评估点"],
        [
            ["Reset", "请执行 reset memory。", "清空记忆，展示 M0。", "可复测性。"],
            ["首次任务", "帮我整理一段关于多模态检索的文献综述。", "输出普通综述。", "无偏好 baseline。"],
            ["反馈", "以后做文献综述时，请按方法类别组织；每篇都标数据集、局限和可复现性；不要夸大结论。", "提取 research_review_method。", "新增 M1：方法分类、数据集/局限/复现、谨慎表述。"],
            ["查看记忆", "展示当前记忆。", "列出研究场景 scope 和证据。", "透明度。"],
            ["再次任务", "帮我综述 RAG 评测方法。", "主动按方法类别组织，并标数据集、局限、可复现性。", "应用 M1，结果质量提升。"],
            ["偏好变化", "如果是头脑风暴研究问题，不要用文献综述模板；只保留谨慎表述。", "将 M1 限定到“文献综述”场景。", "M2：综述模板只在 review 任务生效。"],
            ["第三次任务", "帮我 brainstorm 3 个 RAG 研究问题。", "不再套综述表格，改为问题、假设、验证路径；仍避免夸大。", "旧模板停止影响，新范围生效。"],
            ["删除复测", "删除‘每篇都标可复现性’这条记忆，再做一段简短综述。", "不再强制可复现性字段。", "删除控制。"],
        ],
    )

    heading(doc, "9. 数据模型与接口", 1)
    para(doc, "前端工作台应只依赖 Eval 输出产物，不直接绑定 Skill 内部实现。推荐后端/runner 生成以下 JSON。")
    code(
        doc,
        """{
  "config": {
    "name": "current_module_config",
    "enabled_modules": ["consent_gate", "memory_extraction", "conflict_resolution"]
  },
  "summary": {
    "case_count": 10,
    "config_average": 78.6,
    "effort_reduction": "-62%"
  },
  "cases": [
    {
      "id": "case01",
      "title": "架构方案工作流适应",
      "goal": "验证工作方法提取、应用和条件化更新",
      "score": 84,
      "actions": ["新增", "应用", "更新", "降权"],
      "rounds": [
        {
          "name": "Round 1",
          "title": "记忆形成",
          "user_input": "...",
          "skill_action": "...",
          "memory_action": "新增 workflow_rule",
          "transition": "M0 -> M1",
          "gain": "...",
          "contributes": ["有效记忆提取", "用户控制"],
          "evidence": "..."
        }
      ],
      "memory_journey": [
        {
          "state": "M0",
          "title": "空白记忆",
          "effort_level": "高",
          "effort_score": 92,
          "reason": "用户必须完整说明流程"
        }
      ],
      "ablation": {
        "module": "冲突处理模块",
        "off": ["新旧规则同时 active"],
        "on": ["旧规则降权，新规则条件化生效"]
      },
      "dimension_scores": {
        "reproducibility": 9,
        "memory_extraction": 18,
        "memory_application": 23,
        "update_and_decay": 17,
        "transparency": 8,
        "result_quality": 9
      }
    }
  ]
}""",
    )

    heading(doc, "10. 研发优先级", 1)
    table(
        doc,
        ["优先级", "功能", "验收标准"],
        [
            ["P0", "Case 列表与 Case 详情", "能展示历史 Case、三轮 Round、动作标签和 Case 六维总分。"],
            ["P0", "Round 默认精简展示", "一级页面文字少，细节默认折叠。"],
            ["P0", "Memory 费力度气泡", "高/中/低可解释，规则可展开。"],
            ["P0", "模块 ON/OFF 对比", "每个 Case 能展示一个关键模块的 OFF 后果和 ON 改善。"],
            ["P1", "筛选与排序", "按总分、低分维度、动作类型筛选 Case。"],
            ["P1", "报告导出", "导出 Case 评测报告 Markdown/JSON。"],
            ["P2", "多配置对比", "同时比较多个模块组合的平均分。"],
        ],
    )

    heading(doc, "11. 验收清单", 1)
    bullets(
        doc,
        [
            "页面不再出现 S1/S2/S3 作为主叙事。",
            "Round 卡片没有综合总分。",
            "Case 三轮完成后有六维总分。",
            "Memory State 不打官方分，只展示费力度高/中/低和解释。",
            "每个 Case 至少有一个模块 OFF/ON 对比。",
            "一级页面文案简洁，细节通过展开项查看。",
            "Eval 输出 JSON 可以直接驱动工作台。"
        ],
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
