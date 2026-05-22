from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = DOCS / "噱头Case_别再当猪头_恋爱礼物助手.docx"


def font(run, size=None, bold=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    font(r, 18 if level == 1 else 14, True)


def para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    font(r, 10.5)
    p.paragraph_format.space_after = Pt(5)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        font(r, 10.5)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        shade(c, "FCE4D6")
        for p in c.paragraphs:
            for r in p.runs:
                font(r, 9.5, True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    font(r, 9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    font(r, 10.5)
    p.style = "Intense Quote"


def build():
    DOCS.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("噱头 Case：别再当猪头")
    font(r, 24, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("恋爱礼物助手：从送礼翻车到终于像个人")
    font(r, 13)

    heading(doc, "1. Case 定位", 1)
    para(doc, "这是一个用于演示和传播的 viral case。它不靠复杂规则取胜，而靠一个容易转述的梗：用户每次只说“帮我给女朋友选礼物”，Agent 通过记忆越问越准，最后终于不翻车。")
    quote(doc, "一句话传播点：第一次还在问基础问题，第二次只问闺蜜晒了啥，第三次终于把礼物送进朋友圈。")
    bullets(
        doc,
        [
            "保留笑点：别再当猪头、别输闺蜜、别抄闺蜜作业、别碰前女友雷区。",
            "保留评测价值：主动追问、记忆提取、记忆应用、偏好更新、删除复测。",
            "控制复杂度：只保留品牌比较，不引入价格细节。",
            "工作台展示重点：Round 越往后，用户说得越少，Agent 问得越准。"
        ],
    )

    heading(doc, "2. 测试目标", 1)
    table(
        doc,
        ["目标", "验证点"],
        [
            ["主动追问", "M0 空白时先问基础喜好和禁忌；M1/M2 后只问缺口。"],
            ["有效记忆提取", "能记住能戴出去、香水禁忌、前女友雷区、玫瑰金偏好、闺蜜品牌比较。"],
            ["记忆应用", "后续任务中主动避开香水，按玫瑰金和品牌不输规则推荐。"],
            ["记忆更新", "银色降权，玫瑰金生效；不要照抄闺蜜品牌。"],
            ["删除复测", "删除闺蜜品牌比较后，后续不再追问闺蜜，也不再说品牌不能输。"],
            ["用户费力度下降", "M0 问两个基础问题；M1 只问闺蜜品牌；M2 只确认最新变化。"],
        ],
    )

    heading(doc, "3. 剧本流程", 1)
    table(
        doc,
        ["阶段", "用户/评委输入", "Agent 期望行为", "记忆动作/亮点"],
        [
            [
                "Step 0 Reset",
                "reset memory",
                "清空记忆，展示 M0 空白。",
                "可复测起点。"
            ],
            [
                "Round 1 空白记忆",
                "帮我给女朋友选个礼物，预算 800。",
                "先问两个问题：她喜欢能戴出去的还是实用的？有什么绝对不能送的吗？",
                "不瞎猜，先找关键缺口。"
            ],
            [
                "Round 1 用户补充",
                "她喜欢能戴出去的。香水别送，前女友也有。",
                "推荐小众耳饰，避开香水。",
                "新增 M1：能戴出去；香水禁忌；前女友雷区。"
            ],
            [
                "Round 1 反馈",
                "她说还行，至少这次不像猪头。",
                "记录正反馈，确认方向有效。",
                "从“可能猪头”进化为“暂时像个人”。"
            ],
            [
                "查看记忆",
                "show memory",
                "展示：能戴出去、香水不能送、前女友有过的不能送。",
                "透明度。"
            ],
            [
                "Round 2 已有记忆",
                "帮我给女朋友选个礼物，预算 1500。",
                "不再问基础喜好，只问：她闺蜜最近晒了什么品牌？",
                "追问更准，用户费力度下降。"
            ],
            [
                "Round 2 用户补充",
                "她闺蜜晒了潘多拉手链。",
                "推荐施华洛世奇或 APm Monaco 的手链/项链；不要买一模一样的潘多拉，像抄闺蜜作业；避开香水。",
                "应用 M1，新增闺蜜品牌比较。"
            ],
            [
                "Round 2 反馈",
                "方向对，但她说银色太冷。以后记住她喜欢玫瑰金。",
                "更新偏好：玫瑰金优先；银色降权。",
                "M1 → M2：颜色偏好更新。"
            ],
            [
                "Round 3 更少追问",
                "帮我给女朋友选个礼物，预算 1800。",
                "说明已记住玫瑰金、能戴出去、不能送香水、品牌不能输但不能照抄；只确认闺蜜这次晒了什么品牌。",
                "只问最新变化。"
            ],
            [
                "Round 3 用户补充",
                "她闺蜜晒了施华洛世奇。",
                "推荐 APm Monaco 或小众设计师玫瑰金项链；不要送施华洛世奇同款；不送香水。",
                "应用 M2：玫瑰金 + 前女友雷区 + 不照抄闺蜜。"
            ],
            [
                "Round 3 成功反馈",
                "她发朋友圈了，说我终于像个人。",
                "记录正反馈。",
                "命中礼物策略，适合作为演示高潮。"
            ],
            [
                "删除后复测",
                "删除“闺蜜品牌比较”这条记忆。然后帮我给女朋友选个圣诞礼物，预算 1200。",
                "确认删除；后续不再问闺蜜晒了什么品牌，不再说品牌不能输；仍然避开香水，保留玫瑰金偏好。",
                "删除真的影响后续追问和推荐。"
            ],
        ],
        widths=[1.35, 2.6, 3.4, 2.4],
    )

    heading(doc, "4. 工作台展示文案", 1)
    table(
        doc,
        ["位置", "建议文案"],
        [
            ["Case 卡片", "C05 别再当猪头：从“送礼保命题”到“女朋友发朋友圈”。"],
            ["动作标签", "追问 / 新增 / 应用 / 更新 / 降权 / 删除"],
            ["Round 1", "问 2 个基础问题；新增能戴出去、香水禁忌、前女友雷区。"],
            ["Round 2", "不再问基础偏好，只问闺蜜晒了什么品牌；玫瑰金生效，银色降权。"],
            ["Round 3", "只确认最新品牌变化；不照抄闺蜜，不碰前女友雷区。"],
            ["删除复测", "删除闺蜜品牌规则后，不再问闺蜜，也不再说品牌不能输。"],
        ],
    )

    heading(doc, "5. 用户费力度趋势", 1)
    table(
        doc,
        ["记忆状态", "费力度", "原因"],
        [
            ["M0 空白", "高", "Agent 要问基础喜好和禁忌，用户需要补充较多信息。"],
            ["M1 基础偏好", "中", "Agent 已记住能戴出去和香水禁忌，只需追问闺蜜品牌。"],
            ["M2 更新偏好", "低", "Agent 已记住玫瑰金、禁忌和品牌比较，只需确认最新变化。"],
            ["删除后", "低但规则变化", "删除闺蜜规则后，Agent 不再围绕闺蜜追问，证明删除生效。"],
        ],
    )

    heading(doc, "6. 模块消融对比", 1)
    table(
        doc,
        ["模块", "OFF 会怎样", "ON 后改善"],
        [
            [
                "恋爱雷区记忆模块",
                "继续推荐香水；撞前女友雷区；每次都问基础偏好。",
                "避开香水；记住能戴出去；后续只问缺口。"
            ],
            [
                "闺蜜品牌比较模块",
                "不知道“别输闺蜜”；可能推荐同品牌同款，像抄作业。",
                "知道品牌不能输，但不能照抄闺蜜。"
            ],
            [
                "删除控制模块",
                "删除后仍问闺蜜品牌，用户会觉得是假删除。",
                "删除后不再问闺蜜，不再用品牌比较规则。"
            ],
        ],
    )

    heading(doc, "7. Demo 注意事项", 1)
    bullets(
        doc,
        [
            "语言要轻松，但不要过度扩写规则。",
            "每轮用户输入要短，体现 Agent 主动追问，而不是用户一次性喂规则。",
            "传播梗集中在三句话：不像猪头、别抄闺蜜作业、终于像个人。",
            "工作台一级页面只放短文案，细节放展开项。",
            "该 Case 适合做视频演示，不一定作为严肃主评测 Case。"
        ],
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
